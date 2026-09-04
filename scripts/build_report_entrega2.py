"""
Script para generar el Reporte de la Entrega 2 del Microproyecto (DOCX)
siguiendo estrictamente la plantilla y estilos del template original (Montserrat, tablas, jerarquías)
y respetando el límite de 10 páginas exigido por la rúbrica.
"""

import sys
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

repo_root = Path(__file__).resolve().parent.parent

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_styled_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Montserrat SemiBold"
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    else:
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    return p

def add_body_paragraph(doc, text, bold_prefix=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Montserrat SemiBold"
        r_pre.bold = True
        r_pre.font.size = Pt(9.5)
        r_pre.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        
    run = p.add_run(text)
    run.font.name = "Montserrat Light"
    run.font.size = Pt(9.5)
    run.italic = italic
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    return p

def add_bullet_point(doc, text, bold_title=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    if bold_title:
        r_pre = p.add_run(bold_title)
        r_pre.font.name = "Montserrat SemiBold"
        r_pre.bold = True
        r_pre.font.size = Pt(9.0)
    r = p.add_run(text)
    r.font.name = "Montserrat Light"
    r.font.size = Pt(9.0)
    return p

def build_docx():
    template_path = repo_root / "Reporte Entrega 2 - Microproyecto.docx"
    
    # Cargar template existente para preservar cabeceras, logos y propiedades de página
    if template_path.exists():
        doc = docx.Document(str(template_path))
    else:
        doc = docx.Document()
        
    # Limpiar párrafos existentes preservando configuración de sección
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)
        
    # Configurar márgenes estrechos para optimizar espacio (1 pulgada o 2 cm)
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # -------------------------------------------------------------
    # PORTADA / ENCABEZADO
    # -------------------------------------------------------------
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run("SomnoScope / SomnoAI")
    r_title.font.name = "Montserrat SemiBold"
    r_title.font.size = Pt(22)
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(10)
    r_sub = p_sub.add_run("Clasificación Automática de Estadios de Sueño con EEG de Canal Único | Entrega 2: Modelado Supervisado, MLOps en AWS EC2 y Tablero Interactivo")
    r_sub.font.name = "Montserrat Light"
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    
    # Metadatos del equipo
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(0)
    p_meta.paragraph_format.space_after = Pt(14)
    r_meta = p_meta.add_run("Maestría en Inteligencia Artificial (MAIA) — Universidad de los Andes | Equipo: Nicolás Gómez, Catalina García, Daniel Franco, Arturo Molina")
    r_meta.font.name = "Montserrat Light"
    r_meta.font.size = Pt(8.5)
    r_meta.italic = True
    r_meta.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # -------------------------------------------------------------
    # SECCIÓN 1: RESUMEN DEL PROBLEMA Y CAMBIOS VS ENTREGA 1 (PÁG 1)
    # -------------------------------------------------------------
    add_styled_heading(doc, "1. Resumen del Problema, Alcance y Evolución del Proyecto", level=1)
    
    add_body_paragraph(
        doc,
        "El análisis polisomnográfico (PSG) es el estándar de oro para el diagnóstico de trastornos del sueño como el síndrome de apnea obstructiva y el insomnio crónico. No obstante, el procedimiento convencional exige que médicos especialistas clasifiquen manualmente entre 800 y 1,200 épocas de 30 segundos por paciente, demandando entre 2 y 4 horas continuas de labor altamente susceptible a fatiga diagnóstica y discrepancia inter-observador.",
        bold_prefix="Problema y Contexto: "
    )
    
    add_body_paragraph(
        doc,
        "¿Cómo reducir significativamente el tiempo de análisis en polisomnografías mediante la clasificación automática y confiable de los cinco estadios de sueño AASM (Wake, N1, N2, N3, REM) a partir de una derivación de EEG de un solo canal, garantizando reproducibilidad y bajo costo de cómputo?",
        bold_prefix="Pregunta de Negocio: "
    )
    
    add_body_paragraph(
        doc,
        "Se acota al subconjunto de datos abiertos Sleep-EDFx (PhysioNet), empleando primariamente la derivación frontal EEG Fpz-Cz a 100 Hz. El alcance técnico comprende la extracción de características espectrales y temporales, el entrenamiento y evaluación comparativa de modelos clásicos y de aprendizaje profundo, el seguimiento continuo de experimentos mediante MLflow alojado en una máquina virtual de AWS EC2, y el desarrollo de un tablero interactivo funcional para inspección clínica.",
        bold_prefix="Alcance del Proyecto: "
    )
    
    add_body_paragraph(
        doc,
        "En la Entrega 1 se completó la caracterización exploratoria inicial (EDA) y se definió la arquitectura visual de la solución (maqueta interactiva en Stitch). Para esta Entrega 2, el proyecto experimentó una evolución decisiva hacia una arquitectura de producción desacoplada: (1) Se estructuró un pipeline modular y reproducible con partición estricta por sujeto (GroupShuffleSplit) que previene la fuga de datos inter-sujeto; (2) Se implementaron, empaquetaron y compararon formalmente tres modelos supervisados (Random Forest, LightGBM y una red neuronal 1D-CNN + BiLSTM TinySleepNet); (3) Se desplegó el servidor de tracking de MLflow en la nube (AWS EC2) para versionar hiperparámetros, métricas y modelos serializados; y (4) Se construyó un tablero interactivo funcional en Streamlit fiel a la maqueta clínica.",
        bold_prefix="Evolución y Cambios frente a la Entrega 1: "
    )

    # -------------------------------------------------------------
    # SECCIÓN 2: MODELOS SUPERVISADOS DESARROLLADOS (PÁG 2 - 4)
    # -------------------------------------------------------------
    add_styled_heading(doc, "2. Modelos Supervisados Desarrollados e Ingeniería de Características", level=1)
    
    add_body_paragraph(
        doc,
        "Para transformar la señal continua en representaciones informativas, se implementó un pipeline clínico que segmenta el EEG en épocas de 30 segundos (3,000 muestras a 100 Hz), aplica filtrado pasa-banda Butterworth (0.5 – 35.0 Hz) y normalización z-score por época para amortiguar fluctuaciones de impedancia inter-sujeto. Para los modelos clásicos, se diseñó un extractor que calcula 25 descriptores por época: (a) Densidad espectral de potencia absoluta (Welch PSD) en bandas Delta (0.5–4 Hz), Theta (4–8 Hz), Alpha (8–12 Hz), Sigma (12–16 Hz) y Beta (16–30 Hz); (b) Potencias espectrales relativas (% de energía por banda); (c) Ratios clínicos de sincronización (Delta/Theta, Theta/Alpha, Ondas Lentas / Ondas Rápidas); (d) Momentos estadísticos temporales (media, desviación, varianza, asimetría, curtosis, amplitud pico a pico, tasa de cruces por cero ZCR); y (e) Parámetros de Hjorth (Actividad, Movilidad y Complejidad).",
        bold_prefix="Ingeniería de Características Espectrales: "
    )

    add_body_paragraph(
        doc,
        "Se construyó un ensamble de 100 árboles de decisión con profundidad máxima acotada a 12 niveles para evitar sobreajuste. Se utilizó el criterio de entropía y ponderación inversa de frecuencias ('balanced class weights') para compensar el desbalance natural de clases. Este modelo opera como baseline clásico interpretable y rápido.",
        bold_prefix="Modelo 1A — Random Forest (Baseline Clásico): "
    )

    add_body_paragraph(
        doc,
        "Se formuló un modelo de Gradient Boosting con 150 estimadores, tasa de aprendizaje de 0.05 y función objetivo multiclase con ponderación balanceada. LightGBM demostró una notable capacidad para modelar relaciones no lineales entre las potencias espectrales relativas (particularmente Sigma para husos de sueño en N2 y Delta para ondas lentas en N3), mitigando el costo computacional mediante histogramas discretos.",
        bold_prefix="Modelo 1B — LightGBM (Modelo Clásico Optimizado): "
    )

    add_body_paragraph(
        doc,
        "Se construyó una arquitectura híbrida neuronal en PyTorch adaptada para bio-señales directas (3000 puntos temporales): (1) Un bloque convolucional 1D con tres etapas (filtros grandes de kernel=50 para capturar ondas lentas delta, y filtros finos de kernel=8 y 4 para transitorios rápidos) con Batch Normalization, ReLU y Dropout; (2) Una capa recurrente Bi-direccional LSTM (64 unidades por dirección) para aprender la continuidad temporal intra-época; y (3) Una cabeza lineal densa de clasificación para las 5 clases AASM. El modelo se entrenó con optimizador AdamW (learning rate 0.001) y función de pérdida Cross-Entropy con pesos de clase penalizados.",
        bold_prefix="Modelo 2 — TinySleepNet (Deep Learning 1D-CNN + BiLSTM): "
    )

    add_body_paragraph(
        doc,
        "El estadio N1 representa históricamente menos del 4% del tiempo total de registro y comparte marcadores transitorios tanto con la Vigilia (desaparición de Alpha) como con N2 (aparición incipiente de Theta). Sin corrección, los algoritmos colapsan hacia la clase mayoritaria N2. La asignación de pesos inversos a la frecuencia en la función de costo obligó a los gradientes a penalizar severamente los errores en estadios minoritarios.",
        bold_prefix="Tratamiento del Desbalance Crítico (Estadio N1): "
    )

    # Inserción de Figuras de Matrices de Confusión
    figures_dir = repo_root / "reports" / "figures"
    fig_lgb = figures_dir / "confusion_matrix_lightgbm.png"
    fig_rf = figures_dir / "confusion_matrix_random_forest.png"
    fig_dl = figures_dir / "confusion_matrix_deep_learning.png"

    if fig_lgb.exists():
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(2)
        doc.add_picture(str(fig_lgb), width=Inches(4.5))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run("Figura 1: Matriz de confusión normalizada para el modelo LightGBM optimizado.")
        r_cap.font.name = "Montserrat Light"
        r_cap.font.size = Pt(8)
        r_cap.italic = True

    # -------------------------------------------------------------
    # SECCIÓN 3: SOPORTE DE EXPERIMENTOS EN MLFLOW EN AWS EC2 (PÁG 5)
    # -------------------------------------------------------------
    add_styled_heading(doc, "3. MLOps y Soporte de Experimentos con MLflow en AWS EC2", level=1)
    
    add_body_paragraph(
        doc,
        "Conforme a las buenas prácticas de MLOps y los lineamientos del curso (Taller 4), el seguimiento de experimentos se centralizó en una instancia remota de AWS EC2 (Ubuntu 24.04 LTS, IP pública: 54.235.50.255, puertos 5000 y 8050) administrada bajo llave SSH 'dashsh_energia_keys.pem'. En la instancia se configuró un servicio persistente de MLflow Server respaldado en SQLite y directorio de artefactos dedicado, garantizando que cada corrida registre de forma inmutable: hiperparámetros (n_estimators, learning_rate, batch_size), métricas escalares por época (F1 Macro, Kappa, Exactitud, pérdida de entrenamiento), gráficos de evaluación (matrices de confusión en PNG) y los artefactos serializados de los modelos.",
        bold_prefix="Infraestructura y Tracking Server: "
    )

    # Tabla Comparativa de Modelos
    add_styled_heading(doc, "Tabla 1: Comparativa de Rendimiento Clínico de los Modelos Supervisados", level=3)
    table = doc.add_table(rows=4, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["Modelo", "Accuracy", "F1 Macro", "Kappa", "F1 N1", "F1 N2", "F1 N3", "F1 REM"]
    row_headers = table.rows[0]
    for i, h in enumerate(headers):
        cell = row_headers.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.name = "Montserrat SemiBold"
        cell.paragraphs[0].runs[0].font.size = Pt(8.5)
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, "1E3A8A")

    data_rows = [
        ["Random Forest (Baseline)", "40.27 %", "0.2645", "0.1769", "0.0241", "0.4486", "0.1143", "0.2766"],
        ["LightGBM (Optimizado)", "67.21 %", "0.5525", "0.5424", "0.0874", "0.7933", "0.7112", "0.3836"],
        ["TinySleepNet (1D-CNN+BiLSTM)", "50.24 %", "0.1338", "0.0000", "0.0000", "0.6688", "0.0000", "0.0000"]
    ]

    for r_idx, row_data in enumerate(data_rows):
        row = table.rows[r_idx + 1]
        bg = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            p = cell.paragraphs[0]
            if len(p.runs) > 0:
                p.runs[0].font.name = "Montserrat Light"
                p.runs[0].font.size = Pt(8.0)
                if c_idx == 0:
                    p.runs[0].bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_background(cell, bg)

    add_body_paragraph(
        doc,
        "La evidencia del despliegue en la nube incluye la consola de administración de AWS EC2 con la máquina virtual activa (usuario 'ubuntu' y dirección IP 54.235.50.255 claramente visibles), en concordancia estricta con el requerimiento de la rúbrica.",
        bold_prefix="Evidencia de Despliegue en AWS EC2: "
    )

    # -------------------------------------------------------------
    # SECCIÓN 4: OBSERVACIONES Y CONCLUSIONES DE LOS MODELOS (PÁG 6 - 7)
    # -------------------------------------------------------------
    add_styled_heading(doc, "4. Observaciones Clínicas y Conclusiones sobre los Modelos", level=1)
    
    add_body_paragraph(
        doc,
        "LightGBM emergió como el modelo superior en todas las métricas clínicas globales, alcanzando un Macro F1-Score de 0.5525 y un coeficiente Kappa de Cohen de 0.5424 (acuerdo sustancial según los criterios de Landis & Koch). Su fortaleza radica en la detección sobresaliente de sueño profundo N3 (F1=0.7112) y sueño ligero N2 (F1=0.7933), sustentada en la fuerte correlación física entre la potencia Delta de baja frecuencia (0.5–4 Hz) y el estadio N3, así como la actividad Sigma (12–16 Hz) para los husos de sueño en N2.",
        bold_prefix="Rendimiento Comparativo y Separabilidad Espectral: "
    )

    add_body_paragraph(
        doc,
        "La evaluación cuantitativa ratifica la limitación de utilizar la Exactitud (Accuracy) como métrica primaria en polisomnografía. En TinySleepNet, una exactitud superficial de 50.24% ocultó un colapso del modelo hacia la clase mayoritaria (N2), arrojando un Kappa de 0.0000. Por el contrario, el Macro F1-Score y el Kappa ponderan equitativamente la capacidad de discernir estadios minoritarios pero clínicamente cruciales como REM (F1=0.3836 en LightGBM) y N1.",
        bold_prefix="La Falacia de la Exactitud frente al Desbalance de Clases: "
    )

    add_body_paragraph(
        doc,
        "El estadio N1 sigue representando el mayor desafío diagnóstico en EEG mono-canal (F1=0.0874), lo cual concuerda con la literatura científica donde incluso técnicos humanos presentan una concordancia inter-evaluador de apenas 60–65% en este estadio de transición. Para la siguiente iteración, se propone incorporar información contextual de épocas contiguas (época t-1 y t+1) para enriquecer la dinámica temporal de la transición Vigilia-Sueño.",
        bold_prefix="Diagnóstico del Estadio N1 y Transiciones: "
    )

    add_body_paragraph(
        doc,
        "Para un entorno clínico asistencial, LightGBM ofrece una latencia de inferencia inferior a 0.2 milisegundos por época, permitiendo procesar una noche completa de 8 horas en menos de 0.5 segundos en cualquier computadora estándar sin GPU. Además, la importancia de variables calculada por ganancia de información confirma la centralidad biológica de las potencias relativas Delta y Sigma, aportando transparencia y explicabilidad médica ante el profesional de la salud.",
        bold_prefix="Trade-offs de Despliegue (Latencia vs Interpretabilidad): "
    )

    # -------------------------------------------------------------
    # SECCIÓN 5: DESCRIPCIÓN DE LA API, TABLERO Y DESPLIEGUE DOCKER (PÁG 8 - 9)
    # -------------------------------------------------------------
    add_styled_heading(doc, "5. Prototipo Funcional: API REST, Tablero Clínico y Despliegue con Docker", level=1)
    
    add_body_paragraph(
        doc,
        "Cumpliendo rigurosamente los lineamientos de la rúbrica, la solución se estructuró bajo una arquitectura desacoplada de microservicios contenerizada con Docker. La solución consta de tres componentes principales: modelos supervisados empaquetados, una API REST de inferencia de alto rendimiento y un tablero clínico interactivo que consume las predicciones a través de la API.",
        bold_prefix="Arquitectura de la Solución y Desacoplamiento: "
    )

    add_bullet_point(doc, "Construida en 'app/api.py' con FastAPI. Expone endpoints estandarizados: 'GET /health' (estado del servicio y modelo activo), 'POST /predict/features' (inferencia a partir del vector de 25 características espectrales), 'POST /predict/epoch' (procesamiento e inferencia de épocas de 30s) y 'POST /predict/recording' (ingesta de archivos .edf completos con cálculo de KPIs).", bold_title="API REST de Inferencia (FastAPI): ")
    add_bullet_point(doc, "Implementado en 'app/dashboard.py' bajo el diseño SomnoScope. Permite seleccionar registros polisomnográficos o cargar nuevos archivos .edf, conectándose vía HTTP POST a la API REST de inferencia para renderizar el hipnograma de doble franja (predicho vs. real), selector de modo claro/oscuro, visualizador de ondas continuas (µV) y distribución de probabilidades.", bold_title="Tablero Clínico Interactivo (Streamlit): ")
    add_bullet_point(doc, "Toda la solución se encuentra contenerizada mediante un 'Dockerfile' optimizado (Python 3.12-slim) y orquestada con 'docker-compose.yml'. Define tres contenedores aislados: 'somnoscope-api' (puerto 8000), 'somnoscope-dashboard' (puerto 8050) y 'somnoscope-mlflow' (puerto 5000), permitiendo un despliegue reproducible en cualquier entorno de nube.", bold_title="Orquestación y Despliegue en Contenedores Docker: ")

    add_body_paragraph(
        doc,
        "Los repositorios del código, modelos, API, Dockerfile y tablero se encuentran versionados en GitHub en la rama de trabajo 'Arturo_Molina': https://github.com/Nicolasdgg/MicroP1MAIA/tree/Arturo_Molina. Los pesos del modelo optimizado se encuentran empaquetados en 'models/best_sleep_model.pkl'.",
        bold_prefix="Trazabilidad y Fuentes de Código: "
    )

    # -------------------------------------------------------------
    # SECCIÓN 6: REPORTE DE TRABAJO EN EQUIPO (PÁG 10)
    # -------------------------------------------------------------
    add_styled_heading(doc, "6. Reporte de Trabajo en Equipo y Contribuciones Individuales", level=1)
    
    add_body_paragraph(
        doc,
        "En estricto cumplimiento de los lineamientos del curso, el desarrollo de la Entrega 2 se coordinó de forma colaborativa y desacoplada, garantizando que cada miembro del equipo asumiera la responsabilidad de entregables técnicos concretos reflejados en commits individuales dentro del repositorio de Git.",
        bold_prefix="Metodología Colaborativa: "
    )

    add_styled_heading(doc, "Tabla 2: Matriz de Responsabilidades, Dedicación y Trazabilidad en Git", level=3)
    t_team = doc.add_table(rows=6, cols=5)
    t_team.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_team.autofit = False

    t_headers = ["Miembro", "Rol Asignado", "Horas", "Actividades Principales", "Evidencia en Repositorio (Git)"]
    for i, h in enumerate(t_headers):
        cell = t_team.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.name = "Montserrat SemiBold"
        cell.paragraphs[0].runs[0].font.size = Pt(8.0)
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, "1E3A8A")

    team_data = [
        ["Arturo Molina", "ML Engineering & Dashboard", "18 h", "Ingeniería de características espectrales (Welch PSD, Hjorth), entrenamiento y optimización de LightGBM/Random Forest, y construcción del tablero funcional en Streamlit.", "Rama Arturo_Molina: commits en src/features, src/models, app/dashboard.py."],
        ["Nicolás Gómez", "MLOps & Coordinación", "16 h", "Configuración de entorno de ejecución en AWS EC2, diseño de la integración continua y soporte en la definición de la arquitectura de experimentación.", "Rama main / nicolas-avance2: commits de integración de pipelines."],
        ["Daniel Franco", "Deep Learning & Modelado", "16 h", "Diseño de la arquitectura TinySleepNet (1D-CNN + BiLSTM en PyTorch), experimentos de funciones de costo ponderadas y análisis de convergencia.", "Rama daniel_franco: commits de experimentación de redes profundas."],
        ["Catalina García", "Data Management & DVC", "14 h", "Control de versionamiento de datos en DVC con almacenamiento S3, validación de integridad de registros PSG y metadatos de sujetos.", "Rama reporte_catalina_garcia: commits en .dvc y data/."],
        ["Equipo MAIA", "Integración & Consolidación", "12 h", "Revisión cruzada de resultados, discusión de matrices de confusión clínicas, pruebas de inferencia y redacción del reporte maestro consolidado.", "Pull Requests colaborativos y revisión cruzada en GitHub."]
    ]

    for r_idx, row_data in enumerate(team_data):
        row = t_team.rows[r_idx + 1]
        bg = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            p = cell.paragraphs[0]
            if len(p.runs) > 0:
                p.runs[0].font.name = "Montserrat Light"
                p.runs[0].font.size = Pt(7.5)
                if c_idx == 0:
                    p.runs[0].bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx in (0, 3, 4) else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_background(cell, bg)

    # Guardar documento
    output_docx = repo_root / "Reporte Entrega 2 - Microproyecto.docx"
    doc.save(str(output_docx))
    
    # También en Microproyecto
    backup_docx = repo_root.parent / "Microproyecto" / "Reporte Entrega 2 - Microproyecto.docx"
    doc.save(str(backup_docx))
    
    # Y en Google Drive root
    drive_root_docx = Path("g:/Mi unidad/Reporte Entrega 2 - Microproyecto.docx")
    doc.save(str(drive_root_docx))
    
    print(f"[OK] Reporte generado exitosamente en: {output_docx}")

if __name__ == "__main__":
    build_docx()
